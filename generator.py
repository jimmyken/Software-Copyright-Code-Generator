#!/usr/bin/env python3
# -*- coding: utf-8 -*-
"""
Software Copyright Code Generator
Generates DOCX files compliant with China's software copyright application requirements.
"""

import os
import sys
from pathlib import Path
from docx import Document
from docx.shared import Pt, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
import argparse


class CopyrightCodeGenerator:
    """Generator for software copyright application code documents."""
    
    # Constants for China's software copyright requirements
    LINES_PER_PAGE = 50
    PAGES_FROM_START = 30
    PAGES_FROM_END = 30
    FONT_NAME = 'Courier New'
    FONT_SIZE = 10
    
    # Common source code file extensions
    CODE_EXTENSIONS = {
        '.py', '.java', '.cpp', '.c', '.h', '.hpp', '.cs', '.js', '.ts',
        '.jsx', '.tsx', '.go', '.rs', '.rb', '.php', '.swift', '.kt',
        '.scala', '.m', '.mm', '.sql', '.sh', '.bash', '.pl', '.r',
        '.lua', '.vim', '.asm', '.s', '.f', '.f90', '.pas', '.vb',
        '.xml', '.html', '.css', '.json', '.yaml', '.yml', '.toml',
        '.md', '.txt', '.conf', '.cfg', '.ini', '.properties'
    }
    
    def __init__(self, source_dir, output_file, software_name="软件名称", version="V1.0"):
        """
        Initialize the generator.
        
        Args:
            source_dir: Path to the source code directory
            output_file: Output DOCX file path
            software_name: Name of the software for header
            version: Version of the software
        """
        self.source_dir = Path(source_dir)
        self.output_file = output_file
        self.software_name = software_name
        self.version = version
        self.all_lines = []
        
    def collect_source_files(self):
        """Collect all source code files from the directory."""
        source_files = []
        
        if not self.source_dir.exists():
            raise FileNotFoundError(f"Source directory not found: {self.source_dir}")
        
        for root, dirs, files in os.walk(self.source_dir):
            # Skip common directories that shouldn't be included
            dirs[:] = [d for d in dirs if d not in {
                '.git', '.svn', '.hg', '__pycache__', 'node_modules',
                'venv', 'env', '.env', 'dist', 'build', '.idea',
                '.vscode', 'target', 'bin', 'obj', '.next', 'coverage'
            }]
            
            for file in sorted(files):
                file_path = Path(root) / file
                if file_path.suffix.lower() in self.CODE_EXTENSIONS:
                    source_files.append(file_path)
        
        return sorted(source_files)
    
    def read_source_code(self):
        """Read and collect all source code lines."""
        source_files = self.collect_source_files()
        
        if not source_files:
            raise ValueError("No source code files found in the directory")
        
        for file_path in source_files:
            try:
                # Try different encodings
                content = None
                for encoding in ['utf-8', 'gbk', 'gb2312', 'latin-1']:
                    try:
                        with open(file_path, 'r', encoding=encoding) as f:
                            content = f.read()
                        break
                    except UnicodeDecodeError:
                        continue
                
                if content is None:
                    print(f"Warning: Could not read file {file_path}, skipping...")
                    continue
                
                # Add file header comment
                relative_path = file_path.relative_to(self.source_dir)
                self.all_lines.append(f"// File: {relative_path}")
                self.all_lines.append("")
                
                # Add file content
                lines = content.split('\n')
                self.all_lines.extend(lines)
                self.all_lines.append("")  # Empty line between files
                
            except Exception as e:
                print(f"Warning: Error reading {file_path}: {e}")
                continue
    
    def select_pages(self):
        """
        Select pages according to copyright requirements.
        First 30 pages and last 30 pages, or all if less than 60 pages.
        """
        total_lines = len(self.all_lines)
        total_pages = (total_lines + self.LINES_PER_PAGE - 1) // self.LINES_PER_PAGE
        
        if total_pages <= (self.PAGES_FROM_START + self.PAGES_FROM_END):
            # Return all lines if total pages <= 60
            return self.all_lines
        
        # Select first 30 pages
        first_page_lines = self.PAGES_FROM_START * self.LINES_PER_PAGE
        selected_lines = self.all_lines[:first_page_lines]
        
        # Select last 30 pages
        last_page_lines = self.PAGES_FROM_END * self.LINES_PER_PAGE
        selected_lines.extend(self.all_lines[-last_page_lines:])
        
        return selected_lines
    
    def create_document(self):
        """Create the DOCX document with proper formatting."""
        doc = Document()
        
        # Set document margins (narrow margins)
        sections = doc.sections
        for section in sections:
            section.top_margin = Inches(0.5)
            section.bottom_margin = Inches(0.75)
            section.left_margin = Inches(0.75)
            section.right_margin = Inches(0.75)
            
            # Add header
            header = section.header
            header_para = header.paragraphs[0]
            header_para.text = f"{self.software_name} {self.version}"
            header_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
            header_run = header_para.runs[0]
            header_run.font.size = Pt(10)
            header_run.font.name = self.FONT_NAME
            
            # Add footer with page number
            footer = section.footer
            footer_para = footer.paragraphs[0]
            footer_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
            footer_run = footer_para.add_run()
            footer_run.font.size = Pt(10)
            footer_run.font.name = self.FONT_NAME
            # Add page number field
            fldChar1 = footer_para._element.makeelement(qn('w:fldChar'))
            fldChar1.set(qn('w:fldCharType'), 'begin')
            instrText = footer_para._element.makeelement(qn('w:instrText'))
            instrText.text = "PAGE"
            fldChar2 = footer_para._element.makeelement(qn('w:fldChar'))
            fldChar2.set(qn('w:fldCharType'), 'end')
            footer_run._element.append(fldChar1)
            footer_run._element.append(instrText)
            footer_run._element.append(fldChar2)
        
        # Get selected lines
        selected_lines = self.select_pages()
        
        # Add code lines to document
        for i, line in enumerate(selected_lines):
            # Add page break every LINES_PER_PAGE lines (except for the first page)
            if i > 0 and i % self.LINES_PER_PAGE == 0:
                doc.add_page_break()
            
            # Add line of code
            para = doc.add_paragraph()
            run = para.add_run(line)
            run.font.name = self.FONT_NAME
            run.font.size = Pt(self.FONT_SIZE)
            
            # Set font for both ASCII and East Asian text
            run._element.rPr.rFonts.set(qn('w:eastAsia'), self.FONT_NAME)
            
            # Set line spacing
            para.paragraph_format.line_spacing = 1.0
            para.paragraph_format.space_before = Pt(0)
            para.paragraph_format.space_after = Pt(0)
        
        return doc
    
    def generate(self):
        """Generate the copyright application document."""
        print(f"Collecting source code from: {self.source_dir}")
        self.read_source_code()
        
        total_lines = len(self.all_lines)
        total_pages = (total_lines + self.LINES_PER_PAGE - 1) // self.LINES_PER_PAGE
        print(f"Total lines collected: {total_lines}")
        print(f"Total pages: {total_pages}")
        
        selected_lines = self.select_pages()
        selected_pages = (len(selected_lines) + self.LINES_PER_PAGE - 1) // self.LINES_PER_PAGE
        print(f"Selected pages: {selected_pages}")
        
        print("Creating DOCX document...")
        doc = self.create_document()
        
        print(f"Saving document to: {self.output_file}")
        doc.save(self.output_file)
        print("Done!")
        
        return self.output_file


def main():
    """Main entry point for command-line interface."""
    parser = argparse.ArgumentParser(
        description='Generate DOCX files compliant with China\'s software copyright application requirements.',
        formatter_class=argparse.RawDescriptionHelpFormatter,
        epilog="""
Examples:
  %(prog)s /path/to/source -o output.docx
  %(prog)s /path/to/source -o output.docx -n "MyApp" -v "V2.0"

This tool will:
  1. Scan the source directory for code files
  2. Collect and format the code (50 lines per page)
  3. Generate a DOCX with first 30 and last 30 pages (or all if < 60 pages)
  4. Add proper headers and footers as required by copyright applications
        """
    )
    
    parser.add_argument(
        'source_dir',
        help='Path to the source code directory'
    )
    
    parser.add_argument(
        '-o', '--output',
        default='copyright_code.docx',
        help='Output DOCX file path (default: copyright_code.docx)'
    )
    
    parser.add_argument(
        '-n', '--name',
        default='软件名称',
        help='Software name for header (default: 软件名称)'
    )
    
    parser.add_argument(
        '-v', '--version',
        default='V1.0',
        help='Software version (default: V1.0)'
    )
    
    args = parser.parse_args()
    
    try:
        generator = CopyrightCodeGenerator(
            args.source_dir,
            args.output,
            args.name,
            args.version
        )
        generator.generate()
        return 0
    except Exception as e:
        print(f"Error: {e}", file=sys.stderr)
        import traceback
        traceback.print_exc()
        return 1


if __name__ == '__main__':
    sys.exit(main())
